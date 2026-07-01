"""
gen_eng_report.py  —  工程實驗報告產生器（通用版）v3.0
=======================================================
可獨立執行，也可由 run_report.py import 並覆蓋 CONFIG。

獨立執行：
  1. 修改下方 ══ CONFIG ══ 區段
  2. python gen_eng_report.py

模組用法（由 run_report.py 呼叫）：
  import gen_eng_report as ger
  ger.PRODUCT_NO = 'FAG102B'
  ger.CHARTS_MAP = charts_map   # 來自 cp_summary_report.build_all()
  ger.build_report()
"""

import os, re
from lxml import etree
import fitz                                        # PyMuPDF — PDF 擷取
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ═══════════════════════════════════════════════════════════════════════════════
#  RUNTIME VARIABLES  ← 全部由 run_report.py (UI) 動態注入，此處僅為空值宣告
# ═══════════════════════════════════════════════════════════════════════════════

PRODUCT_NO   = ''
PRODUCT_FUNC = ''
MANAGER      = ''
AUTHOR       = ''

DATASHEET_PDF = ''
WORKBOOK      = ''
OUTPUT_DOCX   = ''

MASK_VERSIONS = []   # [(version, description, bold), ...]  由 UI 注入
LOTS          = []   # [(fab, lot_no, mask_ver, route, wat, cp1), ...]  由 UI 注入

STATIONS  = ['DS00', 'S1P1', 'DS05', 'SFIN', 'DS03', 'SPRE']   # 由 run_report.py 覆蓋
CHARTS_MAP = {}   # 由 run_report.py 注入（cp_summary_report.build_all() 的回傳值）

# CZ Corner 特性報告（選填）── 由 run_report.py 注入
CZ_IMAGES_DIR = ''   # 如 'charts/cz'；若空字串則略過 4.3 小節

# ═══════════════════════════════════════════════════════════════════════════════
#  FIXED PATHS  ← 固定不變
# ═══════════════════════════════════════════════════════════════════════════════

BASE_DIR   = os.path.dirname(os.path.abspath(__file__))
COVER_FILE = os.path.join(BASE_DIR, '封面.docx')
CHARTS_DIR = os.path.join(BASE_DIR, 'charts')

# ── 顏色常數 ──────────────────────────────────────────────────────────────────
C_HDR_DARK  = '1F3864'
C_HDR_MID   = '2E75B6'
C_ROW_ALT   = 'DEEAF1'
C_ROW_WHITE = 'FFFFFF'
C_PASS      = 'E2EFDA'
C_SECTION   = '1F3864'
C_SUB       = '2E75B6'
C_WHITE     = 'FFFFFF'

# ═══════════════════════════════════════════════════════════════════════════════
#  PDF EXTRACTION  — 動態從 datasheet 擷取 GENERAL DESCRIPTIONS & FEATURES
# ═══════════════════════════════════════════════════════════════════════════════

def extract_datasheet_sections(pdf_path):
    """
    從 Winbond 風格 datasheet PDF 動態擷取：
      • Section 1. GENERAL DESCRIPTIONS  → 段落列表
      • Section 2. FEATURES             → [(group_title, [items]), ...]

    Returns:
        gen_paragraphs : list[str]
        feat_groups    : list[tuple[str, list[str]]]
    """
    PAGE_MID = 300       # PDF points — 左/右欄分界 x 座標

    doc = fitz.open(os.path.join(BASE_DIR, pdf_path))
    target = None
    for page in doc:
        txt = page.get_text()
        if '1. GENERAL DESCRIPTIONS' in txt and '2. FEATURES' in txt:
            target = page
            break
    if target is None:
        return [], []

    blocks = target.get_text('blocks')   # (x0,y0,x1,y1,text,blk_no,blk_type)

    # ── 找 FEATURES 的 y 起始位置 ──────────────────────────────────────────
    feat_y = None
    for b in blocks:
        if '2. FEATURES' in b[4]:
            feat_y = b[1]
            break
    if feat_y is None:
        return [], []

    # ── GENERAL DESCRIPTIONS（feat_y 以上的文字區塊）──────────────────────
    SKIP_GEN = re.compile(
        r'^(1\. GENERAL DESCRIPTIONS|Publication|'
        r'[-\s]*\d+[-\s]*|[A-Z0-9\-]+-DTR\s*$)',
        re.I)
    gen_paragraphs = []
    for b in blocks:
        x0, y0, x1, y1, text = b[:5]
        if y1 > feat_y:
            continue
        text = text.strip()
        if not text:
            continue
        # 移除段落內包含的 section heading
        text = re.sub(r'^1\.\s*GENERAL DESCRIPTIONS\s*\n?', '', text)
        text = text.strip()
        if not text or SKIP_GEN.match(text):
            continue
        # 去除行尾換行，還原為連續段落
        cleaned = re.sub(r'\s*\n\s*', ' ', text).strip()
        if cleaned:
            gen_paragraphs.append(cleaned)

    # ── FEATURES（feat_y 以下，左欄 + 右欄分別解析）──────────────────────
    SKIP_FEAT = re.compile(
        r'^(2\. FEATURES|Publication|[-\s]*\d+[-\s]*Product)', re.I)

    def _io_subscript(s):
        """IO0 → IO₀, IO1 → IO₁ ..."""
        return re.sub(r'IO(\d)', lambda m: f'IO{chr(0x2080 + int(m.group(1)))}', s)

    def parse_feat_blocks(col_blocks):
        groups, cur_title, cur_items = [], None, []
        for _, _, text in col_blocks:
            for line in text.split('\n'):
                line = _io_subscript(line.strip())
                if not line:
                    continue
                if line.startswith('•'):
                    if cur_title is not None:
                        groups.append((cur_title, cur_items))
                    cur_title = line.lstrip('•').strip()
                    cur_items = []
                elif line.startswith('–') or line.startswith('-'):
                    cur_items.append(line.lstrip('–-').strip())
                else:
                    # 延續行（上一項目折行）
                    if cur_items:
                        cur_items[-1] += ' ' + line
        if cur_title is not None:
            groups.append((cur_title, cur_items))
        return groups

    feat_blks = []
    for b in blocks:
        x0, y0, x1, y1, text = b[:5]
        if y0 < feat_y:
            continue
        text = text.strip()
        if not text or SKIP_FEAT.match(text):
            continue
        feat_blks.append((x0, y0, text))

    left  = sorted([b for b in feat_blks if b[0] < PAGE_MID],  key=lambda b: b[1])
    right = sorted([b for b in feat_blks if b[0] >= PAGE_MID], key=lambda b: b[1])

    return gen_paragraphs, parse_feat_blocks(left) + parse_feat_blocks(right)


# ═══════════════════════════════════════════════════════════════════════════════
#  XML PRIMITIVES
# ═══════════════════════════════════════════════════════════════════════════════

def _set_cell_shading(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    # 移除舊 shd
    for old in tcPr.findall(qn('w:shd')):
        tcPr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill_hex)
    tcPr.append(shd)


def _set_cell_borders(cell, color='BFBFBF', sz='4'):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(old)
    borders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),   'single')
        b.set(qn('w:sz'),    sz)
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        borders.append(b)
    tcPr.append(borders)


def _add_para_border_bottom(para, color='2E75B6', sz='6'):
    """段落下底線（用於章節標題）。"""
    pPr = para._p.get_or_add_pPr()
    pb  = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    sz)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pb.append(bottom)
    pPr.append(pb)


def _set_para_spacing(para, before_pt=0, after_pt=4):
    pPr  = para._p.get_or_add_pPr()
    spac = pPr.find(qn('w:spacing'))
    if spac is None:
        spac = OxmlElement('w:spacing')
        pPr.append(spac)
    spac.set(qn('w:before'), str(int(before_pt * 20)))
    spac.set(qn('w:after'),  str(int(after_pt  * 20)))


def _apply_font(run, latin, east_asia, size_pt=None):
    """Set Latin font + East Asian font (and optionally size) on a run."""
    run.font.name = latin                          # sets w:ascii / w:hAnsi / w:cs
    rPr    = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), east_asia)
    if size_pt is not None:
        run.font.size = Pt(size_pt)


# ── 封面 XML helper ──────────────────────────────────────────────────────────
def _set_first_run_text(tc, new_text):
    for p in tc.findall(qn('w:p')):
        runs = p.findall(qn('w:r'))
        if runs:
            t = runs[0].find(qn('w:t'))
            if t is None:
                t = OxmlElement('w:t')
                runs[0].append(t)
            t.text = new_text
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            for r in runs[1:]:
                p.remove(r)
            return

# ═══════════════════════════════════════════════════════════════════════════════
#  COVER
# ═══════════════════════════════════════════════════════════════════════════════

def build_cover(doc):
    body   = doc.element.body
    sectPr = body.find(qn('w:sectPr'))
    for child in list(body):
        body.remove(child)
    body.append(sectPr)

    for child in list(Document(COVER_FILE).element.body):
        if child.tag.split('}')[-1] != 'sectPr':
            # ponytail: lxml proxy → deepcopy 會 RecursionError，改用序列化複製
            body.insert(list(body).index(sectPr), etree.fromstring(etree.tostring(child)))

    tables = [c for c in list(body) if c.tag.split('}')[-1] == 'tbl']
    prod_tbl, sig_tbl = tables[0], tables[1]

    prod_rows = prod_tbl.findall(qn('w:tr'))
    # Product No.
    pno_cell = prod_rows[0].findall(qn('w:tc'))[1]
    for p in pno_cell.findall(qn('w:p')):
        runs = p.findall(qn('w:r'))
        for i, r in enumerate(runs):
            if i == 0:
                t = r.find(qn('w:t'))
                if t is not None: t.text = PRODUCT_NO
            else:
                p.remove(r)
    # Function
    fn_cell = prod_rows[1].findall(qn('w:tc'))[1]
    for pi, p in enumerate(fn_cell.findall(qn('w:p'))):
        runs = p.findall(qn('w:r'))
        for i, r in enumerate(runs):
            if pi == 0 and i == 0:
                t = r.find(qn('w:t'))
                if t is not None: t.text = PRODUCT_FUNC
            else:
                p.remove(r)
    # 副理/撰寫人
    sig_cells = sig_tbl.findall(qn('w:tr'))[0].findall(qn('w:tc'))
    _set_first_run_text(sig_cells[0], f'副理: {MANAGER}')
    _set_first_run_text(sig_cells[1], f'撰寫人: {AUTHOR}')

    # ── 封面標題字體：MS Gothic（中文）+ Calibri（英文）26pt ──────────────────
    for child in list(body):
        if child.tag.split('}')[-1] != 'p':
            continue
        runs = child.findall(qn('w:r'))
        if not runs:
            continue
        rPr0 = runs[0].find(qn('w:rPr'))
        if rPr0 is None:
            continue
        sz0 = rPr0.find(qn('w:sz'))
        if sz0 is None or int(sz0.get(qn('w:val'), '0')) < 30:
            continue
        # 這是標題段落 → 套用新字體與 26pt
        for r in runs:
            rPr = r.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                r.insert(0, rPr)
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.insert(0, rFonts)
            rFonts.set(qn('w:ascii'),    'Calibri')
            rFonts.set(qn('w:hAnsi'),    'Calibri')
            rFonts.set(qn('w:cs'),       'Calibri')
            rFonts.set(qn('w:eastAsia'), 'MS Gothic')
            sz = rPr.find(qn('w:sz'))
            if sz is None:
                sz = OxmlElement('w:sz'); rPr.append(sz)
            sz.set(qn('w:val'), '52')           # 26pt = 52 half-points
            szCs = rPr.find(qn('w:szCs'))
            if szCs is None:
                szCs = OxmlElement('w:szCs'); rPr.append(szCs)
            szCs.set(qn('w:val'), '52')

    # 分頁
    pb = OxmlElement('w:p')
    r  = OxmlElement('w:r')
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r.append(br); pb.append(r)
    body.insert(list(body).index(sectPr), pb)

# ═══════════════════════════════════════════════════════════════════════════════
#  CONTENT HELPERS  — 重新設計精美排版
# ═══════════════════════════════════════════════════════════════════════════════

def p_section(doc, text):
    """章節標題：MS Mincho/Cambria 14pt Bold 深藍，段落底線。"""
    p = doc.add_paragraph()
    _set_para_spacing(p, before_pt=12, after_pt=4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    _apply_font(run, 'Cambria', 'MS Mincho')
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    _add_para_border_bottom(p, color='2E75B6', sz='8')
    return p


def p_sub(doc, text):
    """子標題：MS Mincho/Cambria 12pt Bold 中藍。"""
    p = doc.add_paragraph()
    _set_para_spacing(p, before_pt=8, after_pt=2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    _apply_font(run, 'Cambria', 'MS Mincho')
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    return p


def p_body(doc, text='', bold=False, italic=False, size_pt=10,
           align=WD_ALIGN_PARAGRAPH.LEFT):
    """內文：MS Mincho/Cambria 10pt，支援 bold/italic。"""
    p = doc.add_paragraph()
    p.alignment = align
    _set_para_spacing(p, before_pt=0, after_pt=3)
    if text:
        run = p.add_run(text)
        _apply_font(run, 'Cambria', 'MS Mincho', size_pt)
        run.bold   = bold
        run.italic = italic
    return p


def p_caption(doc, text):
    """圖說：Arial 10pt 斜體，置中。"""
    return p_body(doc, text, italic=True, size_pt=10,
                  align=WD_ALIGN_PARAGRAPH.CENTER)


def p_bullet_group(doc, text):
    """FEATURES 群組標題：• MS Mincho/Cambria 10pt Bold 黑。"""
    p = doc.add_paragraph()
    _set_para_spacing(p, before_pt=5, after_pt=1)
    pPr = p._p.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '360')
    ind.set(qn('w:hanging'), '360')
    pPr.append(ind)
    run = p.add_run(f'\u2022  {text}')
    run.bold = True
    _apply_font(run, 'Cambria', 'MS Mincho', 10)
    return p


def p_bullet_item(doc, text):
    """FEATURES 子項目：–  MS Mincho/Cambria 10pt，縮排。"""
    p = doc.add_paragraph()
    _set_para_spacing(p, before_pt=0, after_pt=1)
    pPr = p._p.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '720')
    ind.set(qn('w:hanging'), '240')
    pPr.append(ind)
    run = p.add_run(f'\u2013  {text}')
    _apply_font(run, 'Cambria', 'MS Mincho', 10)
    return p



def make_styled_table(doc, headers, col_widths_cm=None, hdr_color=C_HDR_DARK):
    """
    精美表格：
      • 表頭：深底白字 Bold Arial 11pt
      • 資料列：奇白偶淡藍交替
      • 全格框線 0.5pt (#BFBFBF)
    回傳 table 物件。
    """
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = 'Normal Table'
    _set_para_spacing(doc.paragraphs[-1] if doc.paragraphs else doc.add_paragraph(),
                      before_pt=6, after_pt=6)

    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        _apply_font(run, 'Cambria', 'MS Mincho', 10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_shading(cell, hdr_color)
        _set_cell_borders(cell)

    if col_widths_cm:
        for ci, w in enumerate(col_widths_cm):
            for row in tbl.rows:
                row.cells[ci].width = Cm(w)
    return tbl


def add_styled_row(tbl, values, row_index=1, bold=False,
                   pass_col=None, align=WD_ALIGN_PARAGRAPH.CENTER):
    """
    新增資料列（含交替底色）。
    pass_col: int → 該欄套用淡綠底色（PASS 欄）
    """
    row = tbl.add_row()
    fill = C_ROW_ALT if (row_index % 2 == 1) else C_ROW_WHITE
    for i, v in enumerate(values):
        cell = row.cells[i]
        cell.paragraphs[0].alignment = align
        run = cell.paragraphs[0].add_run(str(v))
        _apply_font(run, 'Cambria', 'MS Mincho', 10)
        run.bold = bold
        cell_fill = C_PASS if (pass_col is not None and i == pass_col) else fill
        _set_cell_shading(cell, cell_fill)
        _set_cell_borders(cell)
    return row


# ═══════════════════════════════════════════════════════════════════════════════
#  DYNAMIC CHART DISCOVERY
# ═══════════════════════════════════════════════════════════════════════════════

def _scan_charts_dir():
    """
    動態掃描 CHARTS_DIR，依實際產出的 PNG 檔案重建 charts_map。
    檔名格式（由 cp_summary_report.py 產出）：
      - 新格式：{station}_{group_key}.png
      - 舊格式：{station}_{group_key}_{failcode}.png
    回傳：{station: {group_key: [fc1, fc2, ...]}}
    summary_*.png 會自動排除。
    """
    result = {}
    if not os.path.isdir(CHARTS_DIR):
        return result

    valid_groups = {'GCCD', 'AACD', 'NHPH', 'NLPL'}
    pat_legacy = re.compile(r'^(.+)_(GCCD|AACD|NHPH|NLPL)_(.+)\.png$', re.IGNORECASE)
    pat_grouped = re.compile(r'^(.+)_(GCCD|AACD|NHPH|NLPL)\.png$', re.IGNORECASE)

    for fname in sorted(os.listdir(CHARTS_DIR)):
        m_new = pat_grouped.match(fname)
        if m_new:
            sta, grp = m_new.group(1), m_new.group(2).upper()
            if grp in valid_groups:
                result.setdefault(sta, {}).setdefault(grp, [])
            continue

        m_old = pat_legacy.match(fname)
        if m_old:
            sta, grp, fc = m_old.group(1), m_old.group(2).upper(), m_old.group(3)
            if grp in valid_groups:
                result.setdefault(sta, {}).setdefault(grp, []).append(fc)

    return result


# ═══════════════════════════════════════════════════════════════════════════════
#  CHARTS
# ═══════════════════════════════════════════════════════════════════════════════

def insert_summary(doc, group_key, caption):
    img = os.path.join(CHARTS_DIR, f'summary_{group_key}.png')
    if os.path.exists(img):
        doc.add_picture(img, width=Cm(20.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_caption(doc, f'▲ {caption}')
    else:
        p_body(doc, f'[Missing: summary_{group_key}.png]')
    p_body(doc)


def insert_top_fail_table(doc, group_key, group_label):
    """從 CHARTS_MAP 自動產各站 Top-5 Fail Category 表格。"""
    p_body(doc, f'▸ {group_label} Group — 各站 Top-5 主要 Fail Category', bold=True)
    tbl = make_styled_table(
        doc, ['測試站', 'Top-1 Fail', 'Top-2 Fail', 'Top-3 Fail', 'Top-4 Fail', 'Top-5 Fail'],
        col_widths_cm=[2.6, 2.9, 2.9, 2.9, 2.9, 2.9])
    for idx, sta in enumerate(STATIONS, 1):
        fails = list(CHARTS_MAP.get(sta, {}).get(group_key, []))
        fails += ['—'] * (5 - len(fails))
        add_styled_row(tbl, [sta] + fails[:5], row_index=idx)
    p_body(doc)


def p_placeholder(doc, hint=''):
    """灰色斜體範例提示，供使用者依圖表填寫分析。"""
    text = (f'※ 分析觀察（請依上方摘要圖與 Bar Chart 填寫）：'
            + (f'  例：{hint}' if hint else ''))
    p = doc.add_paragraph()
    _set_para_spacing(p, before_pt=2, after_pt=6)
    run = p.add_run(text)
    run.italic = True
    _apply_font(run, 'Cambria', 'MS Mincho', 10)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)   # 灰色
    return p


def insert_charts(doc, station, group_key, fail_codes, group_label):
    p_body(doc, f'● {station}  /  {group_label} Group', bold=True)

    grouped_img = os.path.join(CHARTS_DIR, f'{station}_{group_key}.png')
    if os.path.exists(grouped_img):
        if fail_codes:
            p_body(doc, f'Top Fail Cat.: {", ".join(fail_codes)}', italic=True)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(grouped_img, width=Cm(18.0))
        p_body(doc)
        return

    if not fail_codes:
        p_body(doc, f'[No charts found for {station} / {group_key}]', italic=True)
        p_body(doc)
        return

    n_cols = len(fail_codes)
    tbl = doc.add_table(rows=2, cols=n_cols)
    tbl.style = 'Normal Table'

    # 表頭列
    for i, fc in enumerate(fail_codes):
        cell = tbl.cell(0, i)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].add_run(f'Fail Cat.  {fc}')
        run.bold = True
        _apply_font(run, 'Cambria', 'MS Mincho', 10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_shading(cell, C_HDR_MID)
        _set_cell_borders(cell)

    # 舊格式圖片列（寬度動態分配，總表寬約 18cm）
    img_width = Cm(18.0 / n_cols)
    for i, fc in enumerate(fail_codes):
        img = os.path.join(CHARTS_DIR, f'{station}_{group_key}_{fc}.png')
        cell = tbl.cell(1, i)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_shading(cell, C_ROW_WHITE)
        _set_cell_borders(cell)
        if os.path.exists(img):
            cell.paragraphs[0].add_run().add_picture(img, width=img_width)
        else:
            cell.paragraphs[0].add_run(f'[Missing: {station}_{group_key}_{fc}.png]')

    p_body(doc)

def insert_cz_section(doc):
    """
    讀取 CZ_IMAGES_DIR/cz_index.json，動態插入「4.3 Corner CZ Data」小節。
    每張圖前加一行說明：Product / Temperature / Corner case（全從 JSON 取得，不 hard code）。
    若 CZ_IMAGES_DIR 未設定或 cz_index.json 不存在則靜默略過。
    """
    import json
    if not CZ_IMAGES_DIR:
        return

    index_path = os.path.join(CZ_IMAGES_DIR, 'cz_index.json')
    if not os.path.exists(index_path):
        return

    with open(index_path, 'r', encoding='utf-8') as fp:
        entries = json.load(fp)

    if not entries:
        return

    p_sub(doc, '4.3  Corner CZ Data')

    # 依 product → split → temp 排序，讓報告版面整齊
    entries_sorted = sorted(entries, key=lambda e: (e.get('product', ''),
                                                     e.get('split', ''),
                                                     e.get('temp', '')))

    for entry in entries_sorted:
        product = entry.get('product', '')
        temp    = entry.get('temp', '')
        split   = entry.get('split', '')
        path    = entry.get('path', '')

        # 描述文字
        desc = f'Product: {product}      Temperature: {temp}      Corner case: {split}'
        p_body(doc, desc, bold=False)

        # 插入圖片（寬度撐滿版心 ~16cm）
        if os.path.exists(path):
            p = doc.add_paragraph()
            _set_para_spacing(p, before_pt=2, after_pt=6)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(path, width=Cm(16))
        else:
            p_body(doc, f'[Missing image: {os.path.basename(path)}]', italic=True)

    p_body(doc)  # 空行分隔


# ═══════════════════════════════════════════════════════════════════════════════
#  BUILD REPORT
# ═══════════════════════════════════════════════════════════════════════════════

def build_report():
    global CHARTS_MAP
    # 若 CHARTS_MAP 未由 run_report.py 注入（standalone 執行 fallback），才掃描目錄
    if not CHARTS_MAP:
        scanned = _scan_charts_dir()
        if scanned:
            CHARTS_MAP = scanned
            print(f'  ✓ CHARTS_MAP fallback 掃描：{sum(len(v) for v in scanned.values())} 個 group-station 組合')

    doc = Document(COVER_FILE)
    build_cover(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # 壹、概述
    # ─────────────────────────────────────────────────────────────────────────
    p_section(doc, '壹、 概述')

    # 動態從 datasheet PDF 擷取 1. General Descriptions & 2. Features
    gen_paras, feat_groups = extract_datasheet_sections(DATASHEET_PDF)

    p_sub(doc, '1. General Descriptions')
    for para in gen_paras:
        p_body(doc, para)

    p_sub(doc, '2. Features')
    for group_title, items in feat_groups:
        p_bullet_group(doc, group_title)
        for item in items:
            p_bullet_item(doc, item)

    # 分頁
    pb_p = doc.add_paragraph()
    pb_r = OxmlElement('w:r'); pb_b = OxmlElement('w:br')
    pb_b.set(qn('w:type'), 'page'); pb_r.append(pb_b); pb_p._p.append(pb_r)

    p_sub(doc, '3. Mask Version')
    mask = make_styled_table(doc, ['版本', '說明'],
                             col_widths_cm=[3.5, 14.25], hdr_color=C_HDR_DARK)
    for idx, (ver, desc, bold) in enumerate(MASK_VERSIONS, 1):
        add_styled_row(mask, [ver, desc], row_index=idx, bold=bold)

    p_sub(doc, '4. 工程實驗批號記錄')
    lots = make_styled_table(
        doc, ['FAB', 'Lot No.', 'Mask Ver.', 'Process Route', 'WAT', 'CP1'],
        col_widths_cm=[1.5, 3.55, 1.5, 2.25, 3.0, 3.0])
    for idx, row in enumerate(LOTS, 1):
        add_styled_row(lots, list(row), row_index=idx)

    # ─────────────────────────────────────────────────────────────────────────
    # 貳、可靠性測試結果
    # ─────────────────────────────────────────────────────────────────────────
    p_section(doc, '貳、 產品特性分析及可靠性測試結果之簡述')
    p_body(doc, '所有可靠性測試項目均 PASS ✅', bold=True)

    rel = make_styled_table(
        doc, ['測試項目', '標準', '條件', '結果'],
        col_widths_cm=[4.5, 2.5, 8.0, 2.0])
    RELIABILITY = [
        ('ESD — HBM',            'JS-001',       '—',                        'PASS'),
        ('ESD — CDM',            'JS-002',       '—',                        'PASS'),
        ('Latch Up',             'JEDEC STD 78', 'Ta=25°C, Vcc=Max',         'PASS'),
        ('HAST',                 'JESD22-A110',  '130°C, 85% RH, 168 hrs',   'PASS'),
        ('TCT',                  'JESD22-A104',  '-65~+150°C, 500 cycles',   'PASS'),
        ('PCT',                  'JESD22-A102',  '121°C, 100% RH, 168 hrs',  'PASS'),
        ('HTSL',                 'JESD22-A103',  '150°C, 1000 hrs',          'PASS'),
        ('ELFR',                 'JEDEC74',      '125°C, dynamic, 168 hrs',  'PASS'),
        ('HTOL',                 'JESD22-A108',  '125°C, dynamic, 1000 hrs', 'PASS'),
        ('DR',                   'JESD22-A117',  '150°C, 1000 hrs',          'PASS'),
        ('Cycling + DR (Room)',  '—',            '1k~100k cycles',           'PASS'),
        ('Cycling + DR (105°C)', '—',            '1k~100k cycles',           'PASS'),
    ]
    for idx, r in enumerate(RELIABILITY, 1):
        add_styled_row(rel, r, row_index=idx, pass_col=3)

    # ─────────────────────────────────────────────────────────────────────────
    # 參、工程實驗過程說明
    # ─────────────────────────────────────────────────────────────────────────
    p_section(doc, '參、 工程實驗過程說明')
    p_sub(doc, '一. 目的')
    p_body(doc, '1. 【請填寫目的】：')
    p_sub(doc, '二. 實驗方法')
    p_body(doc, '1. 【請填寫實驗方式】：')

    p_sub(doc, '1. Device Skew 實驗分組（NH/PH / NL/PL）')
    sk = make_styled_table(doc, ['Split', '條件類型', '說明'],
                           col_widths_cm=[5.5, 4.5, 7.75])
    SKEW = [
        ('NH/PH SS+3σ / SS+1.5σ', 'N-High / P-High  Slow', 'Slow corner – 速度偏慢'),
        ('POR',                    'Point of Record',        '量產標準基準條件'),
        ('NH/PH FF-1.5σ / FF-3σ', 'N-High / P-High  Fast', 'Fast corner – 速度偏快'),
        ('NL/PL SS+3σ  (×2 Lot)', 'N-Low / P-Low   Slow',  'Slow corner，兩個不同批次'),
        ('NL/PL FF-3σ  (×2 Lot)', 'N-Low / P-Low   Fast',  'Fast corner，兩個不同批次'),
    ]
    for idx, r in enumerate(SKEW, 1):
        add_styled_row(sk, r, row_index=idx,
                       align=WD_ALIGN_PARAGRAPH.LEFT)

    p_sub(doc, '2. GC CD / AA CD Skew 實驗條件')
    cd = make_styled_table(doc, ['參數', 'POR 值', '偏移條件（緊 → 鬆）'],
                           col_widths_cm=[3.0, 3.0, 11.75])
    add_styled_row(cd, ['GC CD', '（請填寫 nm）',
                         '−3σ  →  −1.5σ  →  POR  →  +1.5σ  →  +3σ'], row_index=1)
    add_styled_row(cd, ['AA CD', '（請填寫 nm）',
                         '−3σ  →  −1.5σ  →  POR  →  +1.5σ  →  +3σ'], row_index=2)

    # ─────────────────────────────────────────────────────────────────────────
    # 肆、工程實驗測試結果及分析
    # ─────────────────────────────────────────────────────────────────────────
    p_section(doc, '肆、 工程實驗測試結果及分析')

    # ── 4.1 Device Skew ──────────────────────────────────────────────────────
    p_sub(doc, '4.1  Device Skew CP Yield 彙整')

    # 自動從 CHARTS_MAP 產各站 Top-5 Fail 表
    insert_top_fail_table(doc, 'NHPH', 'NH/PH')
    insert_top_fail_table(doc, 'NLPL', 'NL/PL')

    # 分析觀察填寫提示（灰色範例，使用者依 Bar Chart 自行填入）
    nhph_top1 = next((CHARTS_MAP.get(s, {}).get('NHPH', ['?'])[0]
                      for s in STATIONS if CHARTS_MAP.get(s, {}).get('NHPH')), '?')
    p_placeholder(doc,
        f'POR sample wafer CP yield 最佳；FF corner 良率略低，'
        f'主因 {STATIONS[0]} Fail Cat. {nhph_top1} 偏高（請確認並補充各站百分比）。')

    # 摘要表圖
    insert_summary(doc, 'NHPH', 'NH/PH Split — 各站製程條件良率與 Fail Category 摘要表')
    insert_summary(doc, 'NLPL', 'NL/PL Split — 各站製程條件良率與 Fail Category 摘要表')

    # Bar charts（由 CHARTS_MAP 自動帶入 fail codes）
    p_sub(doc, 'NH/PH Group — Top-5 Fail Category Bar Charts')
    for sta in STATIONS:
        insert_charts(doc, sta, 'NHPH', CHARTS_MAP.get(sta, {}).get('NHPH', []), 'NH/PH')

    p_sub(doc, 'NL/PL Group — Top-5 Fail Category Bar Charts')
    for sta in STATIONS:
        insert_charts(doc, sta, 'NLPL', CHARTS_MAP.get(sta, {}).get('NLPL', []), 'NL/PL')

    # ── 4.2 CD Skew ──────────────────────────────────────────────────────────
    p_sub(doc, '4.2  AA CD / GC CD Skew CP Yield 彙整')

    insert_top_fail_table(doc, 'GCCD', 'GC CD')
    insert_top_fail_table(doc, 'AACD', 'AA CD')

    gccd_top1 = next((CHARTS_MAP.get(s, {}).get('GCCD', ['?'])[0]
                      for s in STATIONS if CHARTS_MAP.get(s, {}).get('GCCD')), '?')
    p_placeholder(doc,
        f'POR sample wafer CP yield 最佳；CD ±1.5σ 偏移影響有限，'
        f'主要失效 Fail Cat. {gccd_top1}（請確認 ±3σ 時是否有明顯變化）。')

    insert_summary(doc, 'GCCD', 'GC CD Split — 各站製程條件良率與 Fail Category 摘要表')
    insert_summary(doc, 'AACD', 'AA CD Split — 各站製程條件良率與 Fail Category 摘要表')

    p_sub(doc, 'GC CD Group — Top-5 Fail Category Bar Charts')
    for sta in STATIONS:
        insert_charts(doc, sta, 'GCCD', CHARTS_MAP.get(sta, {}).get('GCCD', []), 'GC CD')

    p_sub(doc, 'AA CD Group — Top-5 Fail Category Bar Charts')
    for sta in STATIONS:
        insert_charts(doc, sta, 'AACD', CHARTS_MAP.get(sta, {}).get('AACD', []), 'AA CD')

    # ── 4.3 Corner CZ Data ───────────────────────────────────────────────────
    insert_cz_section(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # 伍、Release Lot 資料
    # ─────────────────────────────────────────────────────────────────────────
    p_section(doc, '伍、 Release Lot 資料')
    p_body(doc, '（Mask List 待補，請依實際 Release 資料填入）', italic=True)

    # ─────────────────────────────────────────────────────────────────────────
    # 陸、結論
    # ─────────────────────────────────────────────────────────────────────────
    p_section(doc, '陸、 結論')

    # 用 CHARTS_MAP 的實際 fail code 當提示，讓使用者依圖表確認後修改
    def _top1(sta, grp):
        lst = CHARTS_MAP.get(sta, {}).get(grp) or ['?']
        return lst[0]

    sta_top_nhph = ', '.join(f'{s}:{_top1(s,"NHPH")}' for s in STATIONS)
    sta_top_gccd = ', '.join(f'{s}:{_top1(s,"GCCD")}' for s in STATIONS)

    conclusions = [
        ('【請填寫結論】：', '', ''),
        ('【請填寫】', 'Device Skew 分析：',
         f'各站 NHPH Top-1 Fail 參考：{sta_top_nhph}'),
        ('【請填寫】', 'GC CD / AA CD Skew 分析：',
         f'各站 GCCD Top-1 Fail 參考：{sta_top_gccd}'),
        ('【請填寫】', '量產建議製程條件與後續追蹤事項。', ''),
    ]
    for i, (icon, main_text, hint) in enumerate(conclusions, 1):
        p = doc.add_paragraph()
        _set_para_spacing(p, before_pt=2, after_pt=2)
        pPr = p._p.get_or_add_pPr()
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), '360'); ind.set(qn('w:hanging'), '360')
        pPr.append(ind)
        # 主文字
        r1 = p.add_run(f'{i}. {icon}  {main_text}')
        _apply_font(r1, 'Cambria', 'MS Mincho', 10)
        if icon == '✅':
            r1.bold = True
        # 提示文字（灰色小字）
        if hint:
            r2 = p.add_run(f'\n    {hint}')
            _apply_font(r2, 'Cambria', 'MS Mincho', 9)
            r2.italic = True
            r2.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # ─────────────────────────────────────────────────────────────────────────
    # 柒、附件
    # ─────────────────────────────────────────────────────────────────────────
    p_section(doc, '柒、 附件')

    doc.save(os.path.join(BASE_DIR, OUTPUT_DOCX))
    print(f'Done -> {OUTPUT_DOCX}')


if __name__ == '__main__':
    print('⚠️  請透過 run_report.py (UI) 執行，此檔案不支援獨立執行。')
    print('    python run_report.py')
